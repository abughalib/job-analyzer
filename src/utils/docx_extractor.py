"""DOCX Text Extractor - Extracts text from Microsoft Word documents."""

import logging
from pathlib import Path
from typing import List

from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(docx_path: str | Path) -> str:
    """
    Extract all text content from a DOCX file.

    Args:
        docx_path: Path to the DOCX file.

    Returns:
        Concatenated text extracted from all paragraphs and tables.
    """
    docx_path = Path(docx_path)
    text_chunks: List[str] = []

    try:
        doc = Document(str(docx_path))

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_chunks.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_chunks.append(row_text)

        extracted_text = "\n\n".join(text_chunks).strip()

        if not extracted_text:
            logger.warning(f"No text found in DOCX: {docx_path.name}")
            return "[No extractable text found in this DOCX file.]"

        return extracted_text

    except Exception as e:
        logger.error(
            f"Error extracting text from DOCX {docx_path}: {str(e)}", exc_info=True
        )
        raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")


def save_extracted_text(
    docx_path: str | Path, output_dir: str | Path | None = None
) -> Path:
    """
    Extract text from a DOCX and save it as a .txt file.

    Args:
        docx_path: Path to the source DOCX.
        output_dir: Output directory for the text file. Defaults to the same directory as the DOCX.

    Returns:
        Path to the created text file.
    """
    docx_path = Path(docx_path)
    if not docx_path.is_file():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    output_dir = Path(output_dir) if output_dir else docx_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)

    txt_path = output_dir / f"{docx_path.stem}.txt"

    logger.info(f"Extracting text from: {docx_path}")
    text_content = extract_text_from_docx(docx_path)

    txt_path.write_text(text_content, encoding="utf-8")
    logger.info(f"Extracted text saved to: {txt_path}")

    return txt_path
